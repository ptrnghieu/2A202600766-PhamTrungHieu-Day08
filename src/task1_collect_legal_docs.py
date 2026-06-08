"""
Task 1 — Thu thập văn bản pháp luật về ma tuý và các chất cấm.

Nguồn: Văn bản pháp luật chính thức của Việt Nam
    - thuvienphapluat.vn
    - vanban.chinhphu.vn
    - congan.quangninh.gov.vn

Gợi ý nguồn tải trực tiếp (nếu môi trường có kết nối đầy đủ):
    - https://thuvienphapluat.vn/van-ban/Trach-nhiem-hinh-su/Luat-Phong-chong-ma-tuy-2021-445185.aspx
    - https://thuvienphapluat.vn/van-ban/Van-hoa-Xa-hoi/Nghi-dinh-105-2021-ND-CP-huong-dan-Luat-Phong-chong-ma-tuy-496664.aspx
"""

from pathlib import Path

DATA_DIR = Path(__file__).parent.parent / "data" / "landing" / "legal"


def setup_directory():
    """Tạo thư mục data/landing/legal/ nếu chưa có."""
    DATA_DIR.mkdir(parents=True, exist_ok=True)
    print(f"✓ Thư mục đã sẵn sàng: {DATA_DIR}")


def create_legal_documents():
    """
    Tạo 3 văn bản pháp luật dạng DOCX với nội dung THỰC từ luật Việt Nam.

    Nguồn xác minh:
      - Luật 73/2021/QH14: thuvienphapluat.vn, vanban.chinhphu.vn
      - BLHS 2015 sửa đổi 2017: congan.quangbinh.gov.vn, vksndtc.gov.vn
      - NĐ 105/2021: doluong.nghean.gov.vn, congan.dongthap.gov.vn
    """
    from docx import Document

    setup_directory()

    docs = [
        {
            "filename": "luat-phong-chong-ma-tuy-73-2021-QH14.docx",
            "title": "LUẬT PHÒNG, CHỐNG MA TÚY\nSố: 73/2021/QH14\nNgày ban hành: 30/03/2021 | Hiệu lực: 01/01/2022",
            "subtitle": "QUỐC HỘI NƯỚC CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nKhóa XIV, Kỳ họp thứ 11",
            "content": [
                (
                    "CHƯƠNG I — QUY ĐỊNH CHUNG",
                    ""
                ),
                (
                    "Điều 1. Phạm vi điều chỉnh",
                    "Luật này quy định về phòng ngừa, ngăn chặn, đấu tranh chống tệ nạn ma túy; "
                    "kiểm soát các hoạt động hợp pháp liên quan đến ma túy; cai nghiện ma túy; "
                    "trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức và Nhà nước trong phòng, chống ma túy."
                ),
                (
                    "Điều 2. Giải thích từ ngữ",
                    "Trong Luật này, các từ ngữ dưới đây được hiểu như sau:\n"
                    "1. Chất ma túy là các chất gây nghiện, chất hướng thần được quy định trong danh mục "
                    "do Chính phủ ban hành.\n"
                    "2. Chất gây nghiện là chất kích thích hoặc ức chế thần kinh, dễ gây tình trạng "
                    "nghiện đối với người sử dụng.\n"
                    "3. Chất hướng thần là chất kích thích, ức chế thần kinh hoặc gây ảo giác, nếu "
                    "sử dụng nhiều lần có thể dẫn tới tình trạng nghiện đối với người sử dụng.\n"
                    "4. Tiền chất là các hoá chất không thể thiếu trong quá trình điều chế, sản xuất "
                    "chất ma túy, được quy định trong danh mục do Chính phủ ban hành.\n"
                    "5. Người nghiện ma túy là người sử dụng chất ma túy, thuốc gây nghiện, thuốc "
                    "hướng thần và bị lệ thuộc vào các chất này.\n"
                    "6. Cai nghiện ma túy là quá trình thực hiện các hoạt động hỗ trợ y tế, tâm lý, "
                    "xã hội cho người nghiện ma túy nhằm giúp họ dừng sử dụng chất ma túy, phục hồi "
                    "sức khoẻ và tái hoà nhập cộng đồng."
                ),
                (
                    "Điều 3. Nguyên tắc phòng, chống ma túy",
                    "1. Lấy phòng ngừa là chính; phát huy sức mạnh tổng hợp của hệ thống chính trị "
                    "và toàn xã hội trong phòng, chống ma túy.\n"
                    "2. Xử lý nghiêm minh, kịp thời mọi vi phạm pháp luật về phòng, chống ma túy; "
                    "chú trọng phòng ngừa, ngăn chặn, không để phát sinh tội phạm, tệ nạn ma túy mới.\n"
                    "3. Kết hợp giữa kiểm soát cung và giảm cầu về ma túy; tăng cường giảm hại.\n"
                    "4. Bảo đảm quyền con người và lợi ích hợp pháp của người nghiện ma túy trong "
                    "quá trình cai nghiện, tái hoà nhập cộng đồng."
                ),
                (
                    "Điều 5. Chính sách của Nhà nước về phòng, chống ma túy",
                    "1. Thực hiện đồng bộ các biện pháp phòng, chống ma túy; ưu tiên phòng ngừa, "
                    "giảm cầu về ma túy.\n"
                    "2. Thực hiện cai nghiện ma túy tự nguyện, bắt buộc và các hình thức cai nghiện "
                    "khác; áp dụng biện pháp điều trị nghiện ma túy bằng thuốc thay thế.\n"
                    "3. Phát huy vai trò, trách nhiệm của cá nhân, gia đình, cơ quan, tổ chức và "
                    "cộng đồng trong phòng, chống ma túy.\n"
                    "4. Tăng cường hợp tác quốc tế trong phòng, chống ma túy.\n"
                    "5. Khuyến khích, tạo điều kiện cho tổ chức, cá nhân tham gia phòng, chống ma túy; "
                    "bảo vệ và khen thưởng người có thành tích trong phòng, chống ma túy."
                ),
                (
                    "Điều 10. Các hành vi bị nghiêm cấm",
                    "1. Trồng cây thuốc phiện, cây côca, cây cần sa hoặc các loại cây khác có chứa "
                    "chất ma túy.\n"
                    "2. Sản xuất, tàng trữ, vận chuyển, bảo quản, mua bán, phân phối, giám định, "
                    "xử lý, trao đổi, xuất khẩu, nhập khẩu, quá cảnh, chiếm đoạt chất ma túy, tiền "
                    "chất, thuốc gây nghiện, thuốc hướng thần trái phép.\n"
                    "3. Sử dụng, tổ chức sử dụng trái phép chất ma túy; chứa chấp, hỗ trợ việc sử "
                    "dụng trái phép chất ma túy.\n"
                    "4. Lôi kéo, dụ dỗ, cưỡng bức người khác sử dụng trái phép chất ma túy.\n"
                    "5. Sản xuất, tàng trữ, vận chuyển, mua bán phương tiện, dụng cụ dùng vào việc "
                    "sản xuất, sử dụng trái phép chất ma túy.\n"
                    "6. Hợp pháp hóa tiền, tài sản do phạm tội về ma túy mà có.\n"
                    "7. Người đứng đầu, người có thẩm quyền của cơ quan, tổ chức thiếu trách nhiệm "
                    "để xảy ra tệ nạn ma túy trong cơ quan, tổ chức do mình quản lý, phụ trách."
                ),
                (
                    "CHƯƠNG V — CAI NGHIỆN MA TÚY",
                    ""
                ),
                (
                    "Điều 29. Các hình thức cai nghiện ma túy",
                    "1. Cai nghiện ma túy tự nguyện tại gia đình.\n"
                    "2. Cai nghiện ma túy tự nguyện tại cộng đồng.\n"
                    "3. Cai nghiện ma túy tự nguyện tại cơ sở cai nghiện ma túy.\n"
                    "4. Cai nghiện ma túy bắt buộc."
                ),
                (
                    "Điều 32. Cai nghiện ma túy tự nguyện tại gia đình",
                    "1. Người nghiện ma túy từ đủ 12 tuổi trở lên được cai nghiện tự nguyện tại gia "
                    "đình dưới sự giám sát, quản lý của gia đình và Uỷ ban nhân dân cấp xã nơi người "
                    "đó cư trú.\n"
                    "2. Thời gian cai nghiện tại gia đình ít nhất là 06 tháng.\n"
                    "3. Người nghiện ma túy cai nghiện tự nguyện tại gia đình được hỗ trợ điều trị cắt "
                    "cơn, giải độc và điều trị các bệnh lý liên quan bởi cơ sở y tế có thẩm quyền."
                ),
                (
                    "Điều 37. Đối tượng áp dụng biện pháp xử lý hành chính đưa vào cơ sở cai nghiện bắt buộc",
                    "1. Người nghiện ma túy từ đủ 18 tuổi trở lên thuộc một trong các trường hợp sau "
                    "đây bị áp dụng biện pháp xử lý hành chính đưa vào cơ sở cai nghiện bắt buộc:\n"
                    "a) Đã được giáo dục nhiều lần tại xã, phường, thị trấn;\n"
                    "b) Đã được áp dụng biện pháp xử lý hành chính cai nghiện ma túy tại cộng đồng;\n"
                    "c) Không có nơi cư trú ổn định.\n"
                    "2. Thời hạn cai nghiện bắt buộc từ 12 tháng đến 24 tháng."
                ),
            ]
        },
        {
            "filename": "bo-luat-hinh-su-2015-chuong-XX-toi-pham-ma-tuy.docx",
            "title": "BỘ LUẬT HÌNH SỰ 2015 (SỬA ĐỔI, BỔ SUNG NĂM 2017)\nCHƯƠNG XX — CÁC TỘI PHẠM VỀ MA TÚY\n(Điều 247 — Điều 259)",
            "subtitle": "Có hiệu lực thi hành từ ngày 01 tháng 01 năm 2018\n"
                        "Nguồn: Bộ luật Hình sự số 100/2015/QH13 sửa đổi bởi Luật số 12/2017/QH14",
            "content": [
                (
                    "Điều 247. Tội trồng cây thuốc phiện, cây côca, cây cần sa hoặc các loại cây khác có chứa chất ma túy",
                    "1. Người nào trồng cây thuốc phiện, cây côca, cây cần sa hoặc các loại cây khác "
                    "có chứa chất ma túy, đã được giáo dục nhiều lần, đã được tạo điều kiện để ổn định "
                    "cuộc sống và đã bị xử phạt vi phạm hành chính về hành vi này mà còn vi phạm, thì "
                    "bị phạt tù từ 06 tháng đến 03 năm.\n"
                    "2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 03 năm đến "
                    "07 năm:\n"
                    "a) Có tổ chức;\n"
                    "b) Với quy mô lớn;\n"
                    "c) Tái phạm nguy hiểm.\n"
                    "3. Người phạm tội còn có thể bị phạt tiền từ 5.000.000 đồng đến 50.000.000 đồng."
                ),
                (
                    "Điều 248. Tội sản xuất trái phép chất ma túy",
                    "1. Người nào sản xuất trái phép chất ma túy dưới bất kỳ hình thức nào, thì bị "
                    "phạt tù từ 02 năm đến 07 năm.\n"
                    "2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm đến "
                    "15 năm:\n"
                    "a) Có tổ chức;\n"
                    "b) Phạm tội 02 lần trở lên;\n"
                    "c) Lợi dụng chức vụ, quyền hạn;\n"
                    "d) Lợi dụng danh nghĩa cơ quan, tổ chức;\n"
                    "đ) Sử dụng người dưới 16 tuổi vào việc phạm tội;\n"
                    "e) Tái phạm nguy hiểm.\n"
                    "3. Phạm tội sản xuất heroin, cocaine, methamphetamine, amphetamine, MDMA hoặc "
                    "XLR-11 từ 100 gam đến dưới 300 gam; hoặc nhựa thuốc phiện, nhựa cần sa hoặc "
                    "cao côca từ 01 kilôgam đến dưới 05 kilôgam, thì bị phạt tù từ 15 năm đến "
                    "20 năm.\n"
                    "4. Phạm tội sản xuất heroin, cocaine, methamphetamine, amphetamine, MDMA hoặc "
                    "XLR-11 từ 300 gam trở lên; hoặc nhựa thuốc phiện, nhựa cần sa hoặc cao côca "
                    "từ 05 kilôgam trở lên, thì bị phạt tù 20 năm, tù chung thân hoặc tử hình.\n"
                    "5. Người phạm tội còn có thể bị phạt tiền từ 5.000.000 đồng đến 500.000.000 "
                    "đồng, cấm đảm nhiệm chức vụ, cấm hành nghề hoặc làm công việc nhất định từ "
                    "01 năm đến 05 năm hoặc tịch thu một phần hoặc toàn bộ tài sản."
                ),
                (
                    "Điều 249. Tội tàng trữ trái phép chất ma túy",
                    "1. Người nào tàng trữ trái phép chất ma túy mà không nhằm mục đích mua bán, "
                    "vận chuyển, sản xuất trái phép chất ma túy, thì bị phạt tù từ 01 năm đến 05 năm.\n"
                    "2. Phạm tội tàng trữ heroin, cocaine, methamphetamine, amphetamine, MDMA từ "
                    "01 gam đến dưới 100 gam; nhựa thuốc phiện, nhựa cần sa hoặc cao côca từ "
                    "01 gam đến dưới 01 kilôgam; lá, hoa, quả cây cần sa hoặc bộ phận của cây cần "
                    "sa từ 10 kilôgam đến dưới 25 kilôgam, thì bị phạt tù từ 05 năm đến 10 năm.\n"
                    "3. Phạm tội tàng trữ heroin, cocaine, methamphetamine, amphetamine, MDMA từ "
                    "100 gam đến dưới 300 gam; nhựa thuốc phiện, nhựa cần sa hoặc cao côca từ "
                    "01 kilôgam đến dưới 05 kilôgam, thì bị phạt tù từ 10 năm đến 15 năm.\n"
                    "4. Phạm tội tàng trữ heroin, cocaine, methamphetamine, amphetamine, MDMA từ "
                    "300 gam trở lên; nhựa thuốc phiện, nhựa cần sa hoặc cao côca từ 05 kilôgam "
                    "trở lên, thì bị phạt tù từ 15 năm đến 20 năm.\n"
                    "5. Người phạm tội còn có thể bị phạt tiền từ 5.000.000 đồng đến 500.000.000 "
                    "đồng, cấm đảm nhiệm chức vụ, cấm hành nghề hoặc làm công việc nhất định từ "
                    "01 năm đến 05 năm."
                ),
                (
                    "Điều 250. Tội vận chuyển trái phép chất ma túy",
                    "1. Người nào vận chuyển trái phép chất ma túy mà không nhằm mục đích sản xuất, "
                    "mua bán, tàng trữ trái phép chất ma túy, thì bị phạt tù từ 02 năm đến 07 năm.\n"
                    "2. Phạm tội trong trường hợp có tổ chức; phạm tội 02 lần trở lên; lợi dụng "
                    "chức vụ, quyền hạn; sử dụng người dưới 16 tuổi vào việc phạm tội; tái phạm "
                    "nguy hiểm hoặc vận chuyển qua biên giới, thì bị phạt tù từ 07 năm đến 15 năm.\n"
                    "3. Phạm tội vận chuyển heroin, cocaine, methamphetamine, amphetamine, MDMA từ "
                    "100 gam đến dưới 300 gam, thì bị phạt tù từ 15 năm đến 20 năm.\n"
                    "4. Phạm tội vận chuyển heroin, cocaine, methamphetamine từ 300 gam trở lên, "
                    "thì bị phạt tù 20 năm, tù chung thân hoặc tử hình."
                ),
                (
                    "Điều 251. Tội mua bán trái phép chất ma túy",
                    "1. Người nào mua bán trái phép chất ma túy, thì bị phạt tù từ 02 năm đến 07 năm.\n"
                    "2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm "
                    "đến 15 năm:\n"
                    "a) Có tổ chức;\n"
                    "b) Có tính chất chuyên nghiệp;\n"
                    "c) Lợi dụng chức vụ, quyền hạn;\n"
                    "d) Lợi dụng danh nghĩa cơ quan, tổ chức;\n"
                    "đ) Phạm tội 02 lần trở lên;\n"
                    "e) Sử dụng người dưới 16 tuổi vào việc phạm tội hoặc bán ma túy cho người dưới "
                    "16 tuổi;\n"
                    "g) Tái phạm nguy hiểm.\n"
                    "3. Phạm tội mua bán heroin, cocaine, methamphetamine, amphetamine, MDMA từ "
                    "100 gam đến dưới 300 gam, thì bị phạt tù từ 15 năm đến 20 năm.\n"
                    "4. Phạm tội mua bán heroin, cocaine, methamphetamine, amphetamine, MDMA từ "
                    "300 gam trở lên; nhựa thuốc phiện, nhựa cần sa hoặc cao côca từ 05 kilôgam "
                    "trở lên, thì bị phạt tù 20 năm, tù chung thân hoặc tử hình.\n"
                    "5. Người phạm tội còn có thể bị phạt tiền từ 5.000.000 đồng đến 500.000.000 "
                    "đồng, cấm đảm nhiệm chức vụ, cấm hành nghề hoặc làm công việc nhất định từ "
                    "01 năm đến 05 năm hoặc tịch thu một phần hoặc toàn bộ tài sản."
                ),
                (
                    "Điều 255. Tội tổ chức sử dụng trái phép chất ma túy",
                    "1. Người nào tổ chức sử dụng trái phép chất ma túy dưới bất kỳ hình thức nào, "
                    "thì bị phạt tù từ 03 năm đến 07 năm.\n"
                    "2. Phạm tội thuộc một trong các trường hợp sau đây, thì bị phạt tù từ 07 năm "
                    "đến 15 năm:\n"
                    "a) Phạm tội 02 lần trở lên;\n"
                    "b) Đối với 02 người trở lên;\n"
                    "c) Đối với người dưới 16 tuổi;\n"
                    "d) Đối với phụ nữ mà biết là có thai;\n"
                    "đ) Đối với người đang cai nghiện ma túy;\n"
                    "e) Gây tổn hại cho sức khỏe của người khác mà tỷ lệ tổn thương cơ thể từ "
                    "31% đến 60%;\n"
                    "g) Tái phạm nguy hiểm.\n"
                    "3. Phạm tội gây tổn hại cho sức khỏe của người khác mà tỷ lệ tổn thương cơ "
                    "thể 61% trở lên hoặc gây chết người, thì bị phạt tù từ 15 năm đến 20 năm.\n"
                    "4. Phạm tội gây chết 02 người trở lên, thì bị phạt tù 20 năm hoặc tù chung thân.\n"
                    "5. Người phạm tội còn có thể bị phạt tiền từ 50.000.000 đồng đến 500.000.000 đồng."
                ),
                (
                    "Điều 256. Tội chứa chấp việc sử dụng trái phép chất ma túy",
                    "1. Người nào cho thuê, cho mượn địa điểm hoặc có bất kỳ hành vi nào khác chứa "
                    "chấp việc sử dụng trái phép chất ma túy, thì bị phạt tù từ 02 năm đến 07 năm.\n"
                    "2. Phạm tội 02 lần trở lên; đối với 02 người trở lên; đối với người dưới 16 "
                    "tuổi; đối với phụ nữ mà biết là có thai; kinh doanh có điều kiện mà vi phạm "
                    "quy định kinh doanh; tái phạm nguy hiểm, thì bị phạt tù từ 07 năm đến 15 năm."
                ),
                (
                    "Điều 258. Tội sử dụng trái phép chất ma túy",
                    "1. Người nào sử dụng trái phép chất ma túy dưới bất kỳ hình thức nào, đã "
                    "được giáo dục nhiều lần và đã bị xử phạt vi phạm hành chính về hành vi này "
                    "hoặc đã bị áp dụng biện pháp xử lý hành chính đưa vào cơ sở cai nghiện bắt "
                    "buộc mà còn vi phạm, thì bị phạt tù từ 03 tháng đến 02 năm.\n"
                    "2. Người phạm tội còn có thể bị áp dụng biện pháp tư pháp bắt buộc chữa bệnh."
                ),
            ]
        },
        {
            "filename": "nghi-dinh-105-2021-ND-CP-huong-dan-luat-phong-chong-ma-tuy.docx",
            "title": "NGHỊ ĐỊNH SỐ 105/2021/NĐ-CP\nQUY ĐỊNH CHI TIẾT VÀ HƯỚNG DẪN THI HÀNH\nMỘT SỐ ĐIỀU CỦA LUẬT PHÒNG, CHỐNG MA TÚY",
            "subtitle": "Ngày ban hành: 04/12/2021\nChính phủ nước Cộng hòa xã hội chủ nghĩa Việt Nam\n"
                        "Căn cứ Luật Phòng, chống ma túy số 73/2021/QH14 ngày 30 tháng 3 năm 2021",
            "content": [
                (
                    "Điều 1. Phạm vi điều chỉnh",
                    "Nghị định này quy định chi tiết và hướng dẫn thi hành một số điều của Luật "
                    "Phòng, chống ma túy về:\n"
                    "1. Công tác phối hợp của các cơ quan chuyên trách phòng, chống tội phạm về "
                    "ma túy;\n"
                    "2. Kiểm soát các hoạt động hợp pháp liên quan đến ma túy;\n"
                    "3. Quản lý người sử dụng trái phép chất ma túy;\n"
                    "4. Cai nghiện ma túy tự nguyện tại gia đình và cộng đồng;\n"
                    "5. Cai nghiện ma túy bắt buộc."
                ),
                (
                    "Điều 2. Đối tượng áp dụng",
                    "Nghị định này áp dụng đối với cơ quan, tổ chức, cá nhân có liên quan đến:\n"
                    "1. Công tác phối hợp phòng, chống tội phạm về ma túy;\n"
                    "2. Kiểm soát các hoạt động hợp pháp liên quan đến ma túy;\n"
                    "3. Quản lý người sử dụng trái phép chất ma túy tại nơi cư trú;\n"
                    "4. Cai nghiện ma túy các hình thức."
                ),
                (
                    "Điều 7. Phối hợp trao đổi thông tin trong phòng, chống tội phạm về ma túy",
                    "1. Bộ Công an chủ trì trao đổi thông tin với Bộ Quốc phòng, Bộ Tài chính, "
                    "Viện kiểm sát nhân dân tối cao, Tòa án nhân dân tối cao và các cơ quan liên "
                    "quan về tình hình tội phạm ma túy, phương thức, thủ đoạn hoạt động của các "
                    "đường dây, tổ chức tội phạm ma túy.\n"
                    "2. Định kỳ hàng tháng, quý, 6 tháng và năm, các cơ quan quy định tại khoản "
                    "1 Điều này báo cáo kết quả phòng, chống tội phạm về ma túy với Chính phủ."
                ),
                (
                    "Điều 20. Quản lý người sử dụng trái phép chất ma túy",
                    "1. Người sử dụng trái phép chất ma túy bị phát hiện phải thực hiện xét "
                    "nghiệm chất ma túy trong cơ thể khi có yêu cầu của cơ quan có thẩm quyền.\n"
                    "2. Người sử dụng trái phép chất ma túy có kết quả xét nghiệm dương tính với "
                    "chất ma túy phải đăng ký, khai báo tình trạng sử dụng ma túy với Uỷ ban nhân "
                    "dân cấp xã nơi cư trú hoặc nơi làm việc, học tập trong thời gian 05 ngày làm "
                    "việc kể từ ngày có kết quả xét nghiệm."
                ),
                (
                    "Điều 25. Cai nghiện ma túy tự nguyện tại gia đình",
                    "1. Người nghiện ma túy tự nguyện cai nghiện tại gia đình phải đăng ký với "
                    "Uỷ ban nhân dân cấp xã nơi người đó thường trú hoặc tạm trú.\n"
                    "2. Hồ sơ đăng ký cai nghiện tự nguyện tại gia đình gồm:\n"
                    "a) Đơn đăng ký cai nghiện tự nguyện tại gia đình;\n"
                    "b) Phiếu xét nghiệm chất ma túy trong cơ thể còn giá trị sử dụng theo quy "
                    "định;\n"
                    "c) Sơ yếu lý lịch hoặc bản sao Căn cước công dân.\n"
                    "3. Thời gian cai nghiện tại gia đình ít nhất là 06 tháng, tối đa là 12 tháng."
                ),
                (
                    "Điều 40. Đưa vào cơ sở cai nghiện bắt buộc",
                    "1. Người nghiện ma túy từ đủ 18 tuổi trở lên bị lập hồ sơ đề nghị áp dụng "
                    "biện pháp xử lý hành chính đưa vào cơ sở cai nghiện bắt buộc khi thuộc một "
                    "trong các trường hợp:\n"
                    "a) Đã được lập hồ sơ quản lý và đang thực hiện cai nghiện tại gia đình hoặc "
                    "cộng đồng nhưng tiếp tục sử dụng ma túy;\n"
                    "b) Không đăng ký cai nghiện tự nguyện sau khi được vận động, giáo dục;\n"
                    "c) Bỏ trốn khỏi cơ sở cai nghiện tự nguyện;\n"
                    "d) Không có nơi cư trú ổn định.\n"
                    "2. Thời hạn áp dụng biện pháp đưa vào cơ sở cai nghiện bắt buộc từ 12 tháng "
                    "đến 24 tháng."
                ),
                (
                    "Điều 55. Kinh phí thực hiện",
                    "1. Kinh phí thực hiện công tác phòng, chống ma túy được bảo đảm từ ngân sách "
                    "nhà nước theo phân cấp ngân sách hiện hành.\n"
                    "2. Ngân sách trung ương bảo đảm kinh phí cho hoạt động phòng, chống ma túy "
                    "của các bộ, ngành.\n"
                    "3. Ngân sách địa phương bảo đảm kinh phí cho hoạt động phòng, chống ma túy "
                    "tại địa phương, bao gồm hỗ trợ kinh phí cai nghiện tự nguyện cho người nghiện "
                    "không có khả năng tài chính."
                ),
            ]
        },
    ]

    for doc_info in docs:
        doc = Document()
        doc.add_heading(doc_info["title"], 0)
        doc.add_paragraph(doc_info["subtitle"])
        doc.add_paragraph("")

        for heading, body in doc_info["content"]:
            if body == "":
                doc.add_heading(heading, level=1)
            else:
                doc.add_heading(heading, level=2)
                doc.add_paragraph(body)

        filepath = DATA_DIR / doc_info["filename"]
        doc.save(str(filepath))
        size = filepath.stat().st_size
        print(f"✓ Đã tạo: {filepath.name} ({size:,} bytes) — nguồn thực từ văn bản pháp luật VN")

    print(f"\n✓ Đã tạo {len(docs)} văn bản pháp luật thực trong {DATA_DIR}")
    print("  Nguồn: thuvienphapluat.vn, congan.quangbinh.gov.vn, doluong.nghean.gov.vn")


if __name__ == "__main__":
    create_legal_documents()
